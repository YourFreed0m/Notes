import os
import docx
from docx import Document
import datetime

def openFile(fileName, fileFormat):
    if fileFormat == 'docx':
        doc = Document(r'Notes/{}.docx'.format(fileName))
        all_paras = doc.paragraphs
        for para in all_paras:
            print(para.text)
        
    elif fileFormat == 'txt':
        f = open('Notes/{}.txt'.format(fileName), 'r')
        print(f.read())
        f.close()
        
        
def createFile(fileName, fileFormat):
    if fileFormat == 'docx':
        doc = Document()
        doc.add_heading(str(datetime.date.today()), 0)        
        doc.save(r'Notes/{}.docx'.format(fileName))
        
    else:
        f = open('Notes/{}.{}'.format(fileName, fileFormat), 'w')
        if fileFormat == 'txt':
            f.write(str(datetime.date.today()) + '\n\n')

def changeFile(fileName, fileFormat, text):
    if fileFormat == 'docx':
        doc = Document(r'Notes/{}.docx'.format(fileName))
        doc.add_paragraph(text)
        doc.save(r'Notes/{}.docx'.format(fileName))
        
    elif fileFormat == 'txt':
        f = open('Notes/{}.txt'.format(fileName), 'a')
        f.write(text)
        f.close()
        
def addPicture(fileName, picture):
    doc = Document(r'Notes/{}.docx'.format(fileName))
    doc.add_picture(r"Img/{}".format(picture), width=docx.shared.Inches(5), height=docx.shared.Inches(7))
    doc.save(r'Notes/{}.docx'.format(fileName))
    
    